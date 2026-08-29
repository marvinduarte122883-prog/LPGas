from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = 'FiestaFlow_Project_Summary.docx'
GREEN = '174C3C'
DARK = '18342A'
MUTED = '5D6D66'
PALE = 'EAF2EE'
GOLD = 'A76A18'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true'); trPr.append(el)

def set_fixed_table(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def font(run, size=11, bold=False, color=DARK, italic=False):
    run.font.name = 'Calibri'; run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = RGBColor.from_string(color)

def add_text(doc, text, style=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.10
    font(p.add_run(text))
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.10
    font(p.add_run(text))
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.10
    font(p.add_run(text))
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), size={1:16,2:13,3:12}[level], bold=True, color=GREEN if level < 3 else DARK)
    return p

def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1); table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_table(table, [6.5]); cell = table.cell(0,0); set_cell_shading(cell, PALE)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper() + '  '); font(r, size=10, bold=True, color=GREEN)
    r = p.add_run(body); font(r, size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_table(table, widths)
    header = table.rows[0]; set_repeat_table_header(header)
    for i, h in enumerate(headers):
        set_cell_shading(header.cells[i], 'DDE9E3')
        p = header.cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        font(p.add_run(h), size=9, bold=True, color=GREEN)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
            font(p.add_run(value), size=9.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(11)
for n, size, before, after, color in [('Heading 1',16,16,8,GREEN),('Heading 2',13,12,6,GREEN),('Heading 3',12,8,4,DARK)]:
    s=styles[n]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

# Header and footer
hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(hp.add_run('FIESTAFLOW  |  PROJECT BRIEF'), size=8.5, bold=True, color=MUTED)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run('Confidential working summary • August 2026'), size=8.5, color=MUTED)

# Title block
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(22); p.paragraph_format.space_after=Pt(6)
font(p.add_run('FIESTAFLOW'), size=12, bold=True, color=GREEN)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
font(p.add_run('Petron Fiesta Gas Distribution\nSystem Project Summary'), size=27, bold=True, color=DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(20)
font(p.add_run('A consolidated record of the business model, operating workflows, product requirements, and first-prototype decisions.'), size=12, color=MUTED)
add_callout(doc, 'Purpose', 'This document captures the decisions discussed so far. It is a working business and product brief—not yet a final operating manual, accounting policy, or BIR compliance specification.')

add_heading(doc, '1. Executive overview')
add_text(doc, 'FiestaFlow is being designed as a connected operating system for a new Petron Fiesta Gas distributorship. It will eventually support the main warehouse, company-owned branches, franchisee/wholesale accounts, retail walk-in sales, delivery crews, sales personnel, and central management. The aim is to make cylinder and crate custody, sales, exchanges, payments, delivery proof, vehicle loads, and approvals traceable from the refilling plant through to the customer.')
add_text(doc, 'The current build is a local browser prototype. It demonstrates screens and parts of the POS flow; it does not yet store live business data, authenticate users, issue compliant official receipts, process payments, or provide real GPS tracking.')

add_heading(doc, '2. Business model and operating structure')
add_table(doc, ['Area', 'Agreed direction'], [
    ('Product', 'Petron Fiesta Gas 170g cylinder. Each crate contains 24 cylinders. The business treats cylinders and crates as controlled inventory/custody items.'),
    ('Supply source', 'Stock is picked up from the Petron refilling plant. Empty cylinders and crates are returned to the plant and exchanged for filled stock.'),
    ('Locations', 'A main warehouse, company-owned retail branches, and delivery vehicles. Branches are internal company extensions, not separate entities purchasing stock from head office.'),
    ('Customer groups', 'Franchisees/wholesalers and end users. Franchisees are account-based; end users may be served as normal walk-in retail transactions with minimal personal data.'),
    ('Sales channels', 'Branch POS, warehouse pickup, scheduled delivery, truck sale, and later a franchisee login/customer-facing experience.'),
    ('Branch staffing', 'Initially a small 1–2 person operation, generally a cashier and delivery personnel; no branch manager role is planned at launch.'),
], [1.4,5.1])

add_heading(doc, '3. Inventory, cylinders, crates, and exchange logic')
add_text(doc, 'The physical movement of filled and empty cylinders is the core of the system. Each sale must distinguish between a true exchange and a new issue so truck and branch stock cannot be incorrectly reconciled.')
add_bullet(doc, 'Initial purchase: a new customer buys the cylinder/container and the gas content. The sale therefore includes both the container component and the content component.')
add_bullet(doc, 'Repeat purchase/exchange: where a customer returns the same quantity of empty cylinders, they pay the content/refill price only.')
add_bullet(doc, 'Shortfall: where fewer empties are returned than filled cylinders issued, the unmatched balance is treated as new cylinders and charged accordingly. This prevents a truck sale or new-account sale from being mistakenly recorded as a standard exchange.')
add_bullet(doc, 'More empties than filled items: the transaction should be held for review rather than completed as a normal sale.')
add_bullet(doc, 'Crates: crates are a SKU and must be counted. Operationally, crates are usually exchanged together with cylinders for speed. Some accounts do not purchase crates; the business cannot lend crates to those accounts, so crate custody remains important.')
add_bullet(doc, 'Damaged cylinders: there is no separate damage fee. The customer must purchase a new cylinder if a damaged cylinder needs replacement.')
add_bullet(doc, 'Lost/missing cylinders: the customer may purchase replacement cylinders at the regular cylinder price. The option to carry an outstanding cylinder balance may remain available for future policy decisions, but is not final.')
add_callout(doc, 'Count unit', 'Operations commonly count by crate. The system must still support individual-cylinder sales at branches and show the crate-equivalent where useful.')

add_heading(doc, '4. Pricing and franchise packages')
add_bullet(doc, 'Franchisees are classified by franchise package. Their package determines a fixed peso discount from retail price.')
add_bullet(doc, 'The discount applies to gas content only—not the cylinder/container itself.')
add_bullet(doc, 'The discussed discount range is ₱1–₱3 depending on the franchise package. Pricing should be uniform across company-owned branches.')
add_bullet(doc, 'A newly selected POS prototype price for the Refill item is ₱35. This is an interface-level working value and should be confirmed in the formal price book before live use.')
add_bullet(doc, 'The system must preserve a historical price snapshot per transaction so later changes to a package or price list do not alter past sales.')

add_heading(doc, '5. Accounts, franchisees, and sales commissions')
add_text(doc, 'Franchisee accounts need an internal approval process before regular deliveries can begin. They must submit business information, contact details, and a designated point person. The account file should retain profile information plus purchase history, including what was bought and when.')
add_bullet(doc, 'Truck-sale onboarding: a prospective franchisee may pay and claim their initial inventory during a truck sale. The account remains pending until the truck returns and central office verifies the information.')
add_bullet(doc, 'Pending account pricing: while pending, the truck sale can remain a normal retail-customer sale using the approved price/discount offered under the franchise package.')
add_bullet(doc, 'Sales attribution: the system must record the sales personnel who opened each account.')
add_bullet(doc, 'Commission policy: commissions are expected for opening sales of large accounts (likely ₱20,000 and above). Smaller accounts become eligible after three successful repeat orders. Commission amounts differ by franchise package.')

add_heading(doc, '6. Sales, POS, payments, and official receipts')
add_text(doc, 'The branch POS is expected to serve both end users who buy by cylinder and franchisees who buy by crate or through exchange. The intended checkout process uses physical counts of filled and empty cylinders to determine whether the sale is a new purchase, exchange, or mixed transaction; staff should not have to select a separate “crate exchange” button first.')
add_table(doc, ['POS decision', 'Expected treatment'], [
    ('Filled > 0; empty = 0', 'New sale. Charge gas content plus the required container/cylinder amount.'),
    ('Filled = empty', 'Exchange. Charge content/refill value only and increase empties on hand.'),
    ('Filled > empty > 0', 'Mixed transaction. The matched quantity is an exchange; the remainder is new cylinders.'),
    ('Empty > filled', 'Stop/review. Do not complete as a standard sale until staff explain or correct the count.'),
], [2.0,4.5])
add_bullet(doc, 'Supported payment options: cash, cheque, GCash, and online bank transfer.')
add_bullet(doc, 'For cheque and online/digital payment, the driver or cashier should attach proof. Cheque details and deposit status will later need a dedicated workflow.')
add_bullet(doc, 'The business requires official receipts. Final BIR-compliant invoicing, numbering, issuer details, and printer integration will be defined after the company’s registration details and tax requirements are available.')
add_bullet(doc, 'Cash variance is a major red flag. Any variance—including a claimed miscount—needs a clear reason, central review, and an immutable audit history.')

add_heading(doc, '7. Warehouse, branch, truck, and plant movements')
add_text(doc, 'The stock system must use movements rather than merely editable balances. Each movement should identify the origin, destination, items/quantities, people involved, time, evidence, and approval where applicable.')
add_bullet(doc, 'Receive from Petron plant: record who checked and counted the stock, the driver and vehicle used, and who received it at the destination warehouse or branch.')
add_bullet(doc, 'Return/exchange at plant: record empty cylinders and crates sent back, filled stock received, and any variance.')
add_bullet(doc, 'Warehouse-to-branch replenishment: the main warehouse transfer truck serves both company branches and selected retailer accounts; it is not exclusive to branches.')
add_bullet(doc, 'Truck loading: require a load sheet showing expected filled cylinders/crates, empties, crew, vehicle, and departure/return count. Sales and returns during the route must reconcile to the truck load.')
add_bullet(doc, 'Branch delivery vehicles: each branch has its own drivers and vehicles, focused on surrounding retailers. Vehicle roster and load capacity must be adjustable as more vehicles are added.')
add_bullet(doc, 'Known capacity: the main truck is rated for 1,672 crates (40,128 cylinders) at any time. Other vehicle capacities will be configured later.')
add_bullet(doc, 'Daily counts are preferred. A count captures the expected book balance at the count time, the actual physical count, the counter, and the resulting variance. Daily changes from sales and replenishment are explained through recorded movements, not ignored.')

add_heading(doc, '8. Chain of custody and traceability')
add_text(doc, 'The business wants strong custody controls without slowing day-to-day operations unnecessarily. Each 170g cylinder has a QR seal, although it is not yet confirmed whether the code is unique to a cylinder or a production batch. The initial approach should support scanning and later determine the appropriate uniqueness level after supplier confirmation.')
add_bullet(doc, 'Phone scanning: drivers can use a phone camera to scan supported QR/barcodes as part of truck loading, delivery, return, or count workflows. The system can generate barcodes for internally managed items as needed.')
add_bullet(doc, 'Crate tagging: because crates are swapped at the Petron plant, retagging every crate on every cycle could slow operations. The agreed direction is to keep crate custody/count controls first and avoid mandatory individual crate tagging at launch unless loss patterns justify it.')
add_bullet(doc, 'Chain of custody: retain an event history showing who received, counted, loaded, delivered, transferred, returned, adjusted, or approved each inventory movement.')
add_bullet(doc, 'Audit trail: overrides and changes must retain who made the change, when, what was changed, the prior and new values, a reason, and approval status where required. The history should not be editable by regular users.')

add_heading(doc, '9. Delivery, route, and driver workflow')
add_text(doc, 'Drivers work with at least one assistant and/or salesman. The mobile workflow must be practical for delivery operations and continue to collect data if signal is weak.')
add_bullet(doc, 'Daily route assignment: central management assigns drivers, assistants/salesmen, the vehicle, planned stops, load, and route sequence.')
add_bullet(doc, 'Delivery proof: capture the delivered and returned quantities, payment method, proof where needed, recipient signature, delivery photo, and geotag/location when available.')
add_bullet(doc, 'Exceptions: drivers must record damaged cylinders, failed deliveries, partial delivery, payment issues, and stock discrepancies. These should feed route reconciliation and approvals.')
add_bullet(doc, 'Unplanned acquisition stop: the driver can request a stop to acquire a new account. It remains pending central approval; approved exceptions should be reflected in the route/geofence record.')
add_bullet(doc, 'Offline-first behavior: the device should timestamp the local action time, retain the event while offline, and sync when connectivity returns. The system should keep both device-captured and server-received timestamps where relevant.')
add_bullet(doc, 'Route reconciliation: on return, the truck should reconcile loaded stock + items received/returned + delivered/sold stock + actual count + cash/payment proofs.')

add_heading(doc, '10. Attendance, live location, communication, and geofencing')
add_bullet(doc, 'Employees should be able to time in/out through the app. Attendance can be validated using a location/geofence around the assigned warehouse, branch, or starting location, with a photo/selfie or supervisor review added later if needed.')
add_bullet(doc, 'Live employee location is intended only during working hours, with clear policy, consent, and role-based access. It should not be presented as all-hours personal tracking.')
add_bullet(doc, 'Truck route geofencing is desired: vehicles should remain on their assigned service route except for exceptions approved by central office. Alerts should be designed to distinguish legitimate detours and loss of GPS/signal from genuine deviations.')
add_bullet(doc, 'Radios are feasible for dispatch voice communication within a roughly 20 km area only after a site survey, local licensing/approved equipment check, terrain review, and repeaters if needed. They are a useful backup but not a replacement for the app’s offline record keeping.')
add_bullet(doc, 'Poor signal areas are expected. The final driver app should cache route, records, signatures, photos, and drafts locally, then synchronize safely later.')

add_heading(doc, '11. Roles and access model')
add_table(doc, ['Role', 'Initial responsibility'], [
    ('Central management / owner', 'Approvals, price and package controls, central reporting, audit review, stock and cash exceptions.'),
    ('Warehouse manager / receiver', 'Plant receiving, counting, loading, warehouse custody, and transfers.'),
    ('Cashier', 'Branch POS, payment capture, receipt workflow, and branch daily count.'),
    ('Driver', 'Assigned route, delivery completion, proof, payment record, exceptions, and return reconciliation.'),
    ('Assistant / salesman on truck', 'Can assist with delivery counts, proof, and account acquisition under assigned permissions.'),
    ('Sales personnel', 'Prospect/account opening and sales attribution; commission visibility based on approved policy.'),
    ('Franchisee user (future)', 'Account profile, order/history view, and possible ordering/log-in functionality.'),
], [1.65,4.85])
add_text(doc, 'The role structure must be configurable so new roles, supervisors, managers, or approval layers can be added without rebuilding the system.')

add_heading(doc, '12. Customer experience and future customer application')
add_bullet(doc, 'A separate end-user application is not needed at launch. End users can be handled through branch POS and delivery/truck workflows.')
add_bullet(doc, 'A future franchisee login can be integrated into the same overall platform. Potential features include account profile, purchase history, order request, payment status, and delivery information.')
add_bullet(doc, 'Central management and, eventually, end users may see active franchisee locations. Public-facing location visibility needs a clear business rule, accuracy policy, and consent/marketing treatment before release.')
add_bullet(doc, 'Low-stock control: the desired baseline is stock equivalent to three times current sales volume at the main warehouse and each branch. The exact time window and calculation method (for example, 3× daily average sales) should be confirmed in the inventory policy.')

add_heading(doc, '13. Technical direction, privacy, and long-term considerations')
add_bullet(doc, 'The static prototype can use browser local storage only for demonstrations. It is not appropriate as the real operational database because it is device-specific, can be cleared, and cannot safely coordinate multiple people, branches, and trucks.')
add_bullet(doc, 'A hosted backend such as Supabase is the likely next step when converting the prototype into a working system. It can provide a shared database, user sign-in, permissions, file storage, and server-side logic. A paid plan may eventually be needed based on users, storage, traffic, and support requirements, but it is not necessary before the product flow is finalized.')
add_bullet(doc, 'Photos, signatures, payment evidence, and attendance evidence should be stored in object/file storage rather than the main transaction database. Use compression, retention rules, secure access, and links from the transaction records to control long-term cost.')
add_bullet(doc, 'Live maps, GPS, payment gateway, printer services, invoice compliance, and cloud hosting are later integrations. They require registered business details, provider accounts, credentials, accepted terms, and production policies.')
add_bullet(doc, 'An internal system clock should be authoritative for official records. Device time may still be retained as supporting context for offline events, but server time should control approval, audit, and receipt records after synchronization.')

add_heading(doc, '14. Current prototype status')
add_text(doc, 'A local HTML/CSS/JavaScript prototype has been created in the project folder. It is a visual and interaction prototype; it does not yet use Supabase or a server.')
add_table(doc, ['Prototype area', 'Current state'], [
    ('Main navigation', 'Basic static pages for overview, POS, dispatch, inventory, accounts, approvals, vehicles, reports, and driver route view.'),
    ('POS layout', 'Customer Type is in the left column; Item is in the right column; Payment Method is a single row beneath them.'),
    ('POS item choices', 'Refill, Cylinder, and Crate (24). Refill was set to a working ₱35 price. Filled/empty quantities use a custom on-screen number pad and start blank.'),
    ('Transaction inference', 'Prototype calculates new sale, exchange, mixed sale, or count review using filled and empty counts.'),
    ('Payment proof', 'Prototype requests proof for non-cash payment choices, but it does not upload or retain files yet.'),
    ('Data and sign-in', 'Not implemented. Sample data only; no live multi-user data or permissions.'),
], [1.75,4.75])

add_heading(doc, '15. Recommended delivery plan')
add_number(doc, 'Finish the prototype flows. Review every screen and ensure each button, decision, and exception is understandable before backend work begins.')
add_number(doc, 'Define the operational master data. Confirm product price book, packages/discounts, branch list, vehicle capacities, roles, approval thresholds, receipt requirements, and transaction naming.')
add_number(doc, 'Build the backend foundation. Set up the shared database, user sign-in, role permissions, locations, products, inventory balances, and audit records.')
add_number(doc, 'Make core workflows real. Persist POS sales/exchanges, stock transfers, truck loading, delivery records, daily counts, payments, and approvals.')
add_number(doc, 'Build driver operations. Add mobile sign-in, route assignment, offline queue, proofs, payment evidence, exception capture, and route reconciliation.')
add_number(doc, 'Add external services. Integrate cloud hosting, maps/GPS, payment providers, printers, and BIR-compliant invoicing only once the required accounts and registrations are ready.')
add_number(doc, 'Pilot deliberately. Start with one warehouse, one branch, and one truck using test data; then use monitored real transactions after the team is comfortable.')

add_heading(doc, '16. Decisions still to confirm')
add_bullet(doc, 'Whether each cylinder QR code is unique to an individual cylinder or identifies a batch; confirm with Petron before designing mandatory scanning rules.')
add_bullet(doc, 'Formal retail price book: content/refill prices, cylinder/container price, crate price, franchise package rules, taxes, and effective dates.')
add_bullet(doc, 'Whether “outstanding cylinder balance” will be permitted, how it is settled, and who can approve it.')
add_bullet(doc, 'Exact definition of a successful repeat order for commission purposes and commission amounts by franchise package.')
add_bullet(doc, 'BIR registration, invoice/official-receipt format, taxpayer details, VAT status, and approved printer/e-invoicing process.')
add_bullet(doc, 'Mobile device approach, employee privacy policy, phone data plans, GPS rules, photo retention duration, and radio equipment/site survey.')
add_bullet(doc, 'Whether franchisee locations will be public to end users and what “active” means for map display.')
add_bullet(doc, 'Low-stock formula: exact sales period, buffer definitions, and alert recipients for warehouse and branch replenishment.')

add_callout(doc, 'Next practical action', 'Continue refining the POS and other prototype workflows until the team agrees on the operational flow. Once that flow is settled, set up the real backend and convert the prototype screens one workflow at a time.')

add_heading(doc, '17. Platform, screens, and device approach')
add_text(doc, 'The intended solution is one connected platform with role-based screens rather than completely unrelated applications. Central management will use a wider dashboard for approvals, inventory, accounts, routes, reports, and configuration. Cashiers will use a focused branch POS. Drivers, assistants, and truck sales personnel will use a mobile-friendly field workflow. A future franchisee login can use the same shared platform and data model.')
add_bullet(doc, 'The central/management experience and driver experience may look different, but they should read and write to the same controlled system so branch, warehouse, truck, and customer records remain consistent.')
add_bullet(doc, 'A branch POS terminal can be a tablet or a phone when the interface is designed with large touch targets, camera support, and a compatible receipt printer. A tablet is generally more comfortable for a cashier; a phone is practical for drivers and backup use.')
add_bullet(doc, 'The user interface should be installable as a web app (PWA) later, so it can behave like an app on phones and tablets without requiring a separate native iOS/Android build at the first stage.')
add_bullet(doc, 'Invoice/receipt printers can later connect through a compatible Bluetooth, Wi-Fi, network, or USB printing path. The final choice depends on the printer model, device operating system, and BIR-compliant receipt process.')
add_bullet(doc, 'All screens should work from one shared set of permissions: a cashier should not see or approve central controls; a driver should see only assigned work; central management can review all required records.')

add_heading(doc, '18. Core data model to be built')
add_text(doc, 'Before live transactions begin, the backend should be built around the records below. These records make it possible to calculate balances from controlled movements, preserve history, and limit access by role and location.')
add_table(doc, ['Record group', 'Key information required'], [
    ('Users and roles', 'Employee profile, role, active status, assigned location/vehicle where relevant, sign-in credentials, permissions, and audit identity.'),
    ('Locations', 'Main warehouse, company-owned branches, possible pickup points, and their stock/low-stock settings.'),
    ('Products and price book', '170g cylinder, content/refill, crate SKU, container component, base price, package discounts, taxes, effective dates, and price history.'),
    ('Franchise packages and accounts', 'Package, discount, business profile, contacts/point person, sales owner, approval status, opening sale, repeat order count, and commission status.'),
    ('Inventory and movements', 'Filled/empty cylinders and crates by location/vehicle, transfer/plant/sale/count movement type, quantities, source/destination, counts, reasons, and people involved.'),
    ('Vehicles and routes', 'Vehicle roster, capacity, home location, crew, assignments, planned/actual stops, GPS events, approved route exceptions, and route completion.'),
    ('Sales and payments', 'Transaction lines, inferred exchange/new/mixed classification, payment method, proof attachment, totals, price snapshot, receipt status, and cashier/driver.'),
    ('Evidence and approvals', 'Photos, signatures, files, geotags, timestamps, approval requests/decisions, exception reasons, and immutable audit events.'),
], [1.7,4.8])
add_text(doc, 'A preliminary Supabase/Postgres schema was drafted in the project folder as a planning aid. It is not connected to the prototype and should be reviewed against the final workflows before being used in production.')

add_heading(doc, '19. Detailed prototype and POS interface decisions')
add_text(doc, 'The local prototype was intentionally iterated before backend development so the checkout flow can be comfortable for branch personnel. The following interface decisions are currently reflected or planned in the POS.')
add_bullet(doc, 'The former top-level sale-type buttons (for example, a separate crate-exchange button) were removed from the intended flow. The system should infer the stock movement from filled and empty counts rather than asking staff to classify it first.')
add_bullet(doc, 'Customer Type is a left-side choice group. Item is a right-side choice group. Payment Method sits below as a separate full-width row.')
add_bullet(doc, 'Customer Type choices currently include Walk-in end user and Franchisee. Franchisee selection will ultimately retrieve the approved account, package, and applicable content discount.')
add_bullet(doc, 'Item choices are ordered Refill, Cylinder, then Crate (24). Refill is a selectable option and uses the current prototype working price of ₱35.')
add_bullet(doc, 'Payment choices are presented as color-coded touch boxes in one line: Cash, GCash, Bank transfer, and Cheque. The sizing was adjusted to keep labels visible while remaining easy to tap.')
add_bullet(doc, 'Filled and empty quantity fields start blank. Selecting a field opens an on-screen number pad instead of relying on small browser up/down controls. This supports faster tablet/phone use.')
add_bullet(doc, 'The receipt panel updates the description, total, stock note, and count-review warning from the entered figures. In the current prototype this is illustrative only and not yet an official invoice.')
add_bullet(doc, 'The browser version can be opened locally by opening index.html. It is appropriate for discussion and design review, not real operations or shared data entry.')

add_heading(doc, '20. Infrastructure, performance, and storage decisions')
add_text(doc, 'The expected solution remains viable even with live location, maps, offline field use, photo evidence, and several locations, provided it is designed as a normal cloud application with the right separation of responsibilities.')
add_bullet(doc, 'Driver device RAM: the system does not need to load every customer, map, photo, or stock record into a phone at once. The driver app should fetch/cache only the assigned route, current load, necessary account details, and pending offline records. This keeps the device experience practical.')
add_bullet(doc, 'Server capacity: live GPS should be recorded at a sensible interval and only during approved working/route time. Maps can be loaded on demand. Photos and signatures belong in file storage, not in database rows. These practices keep compute and database load manageable.')
add_bullet(doc, 'Photo storage: delivery proof, payment proof, attendance evidence, and other photos can grow over time. The future policy should define image compression, maximum upload size, attachment types, retention duration, archiving/deletion rules, and access control.')
add_bullet(doc, 'Offline conflict handling: each offline action requires a unique record ID, locally captured time, user/device identity, and later server acknowledgement. The server should reject duplicate submissions and flag conflicting stock/cash actions for review rather than silently overwriting data.')
add_bullet(doc, 'Security: use individual user accounts, strong password/reset flows, role-based access, secure file URLs, encrypted transport, audit logging, and regular backups. Do not use a shared cashier or driver password for production operations.')

add_heading(doc, '21. External services and prerequisites')
add_table(doc, ['Service', 'What is needed before live integration'], [
    ('Payment gateway / collection', 'Registered business details, chosen provider(s), merchant onboarding, settlement account, API credentials, refund/exception process, and finance reconciliation rules.'),
    ('Maps and live GPS', 'Selected map provider, API key/billing account, privacy notice, employee consent/process, location update intervals, geofence rules, and offline fallback.'),
    ('Official receipts / invoicing', 'Business registration, BIR requirements, taxpayer/branch information, receipt sequencing, VAT/tax configuration, approval process, and selected compliant invoice/printing solution.'),
    ('Receipt printer', 'Compatible printer model(s), tablet/phone operating system, connection method, paper format, and tested print layout.'),
    ('Cloud hosting', 'Production database/storage account, domain if desired, backups, monitoring, access administration, and support/incident process.'),
], [1.7,4.8])

add_heading(doc, '22. Project working history and safeguards')
add_bullet(doc, 'The project began as a static website prototype so the business could clarify how it should work before committing to backend services or ongoing hosting costs.')
add_bullet(doc, 'Git was initialized for version control. Future meaningful changes should be committed with short descriptions so earlier working versions can be recovered safely. Avoid using destructive Git commands unless the exact effect is understood.')
add_bullet(doc, 'The project currently contains the prototype front-end files, preliminary backend-planning files, and this summary document. Changes to the prototype do not automatically create a working database, a mobile app-store app, or a production deployment.')
add_bullet(doc, 'The best next design discipline is to settle one workflow at a time—starting with branch POS—then write acceptance rules for normal, exception, and approval cases before implementing it against the real backend.')

add_heading(doc, '23. Comprehensive launch checklist')
add_bullet(doc, 'Confirm all product, crate, cylinder/container, content/refill, and franchise-package prices, including effective dates and tax treatment.')
add_bullet(doc, 'Confirm customer onboarding requirements, pending-account rules, commission conditions, and who may approve exceptions.')
add_bullet(doc, 'Confirm warehouse/branch list, vehicle capacities, daily count process, route process, and the exact truck return/reconciliation checklist.')
add_bullet(doc, 'Obtain Petron clarification on QR/seal coding and decide which items require scan-level traceability at launch.')
add_bullet(doc, 'Confirm official-receipt/BIR requirements after the business registration is complete.')
add_bullet(doc, 'Write employee location, attendance, evidence-photo, data retention, and radio-use policies before live tracking begins.')
add_bullet(doc, 'Finish and test the prototype flows with the intended users: owner/central management, cashier, driver, assistant, warehouse receiver, and sales personnel.')
add_bullet(doc, 'Set up the shared backend only after the initial workflows and master data are approved, then pilot with one warehouse, one branch, and one truck.')

doc.save(OUT)
print(OUT)
